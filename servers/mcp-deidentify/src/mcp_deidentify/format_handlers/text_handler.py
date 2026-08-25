"""Text and DOCX de-identification handler for mcp-deidentify.

Supports:
  source_format="txt"  -- de-identify a plain text string
  source_format="docx" -- de-identify a DOCX file and write a new file to output_path

DRY_RUN=true: no Haiku calls, no disk writes; returns synthetic fixture output.
"""

import asyncio
import logging
import os
from pathlib import Path

from mcp_deidentify import config
from mcp_deidentify.engine import extract_entities, replace_entities

logger = logging.getLogger(__name__)

OUTPUT_DIR = os.getenv("DEIDENTIFY_OUTPUT_DIR", "data/patients")

# Synthetic de-identified text returned in DRY_RUN mode
_SYNTHETIC_DEID_NOTE = (
    "Patient: PAT-NAME-001 | Physician: Dr. ONC-001 | "
    "Facility: FAC-001 | MRN: MRN-REDACTED-001 | "
    "Accession: ACCESSION-001 | DOB: DOB-REDACTED"
)


def _default_output_path(patient_id: str, source_path: str) -> str:
    """Build the default DOCX output path from patient_id and source file name."""
    basename = Path(source_path).stem
    out_dir = Path(OUTPUT_DIR) / patient_id / "deidentified"
    return str(out_dir / f"{basename}_deid.docx")


async def deidentify_text_string(
    text: str,
    patient_id: str,
    session_key: dict,
) -> tuple[str, list[dict]]:
    """De-identify a plain text string.

    Args:
        text:        Input text containing potential PII.
        patient_id:  Patient identifier for code generation.
        session_key: Mutable anonymization key dict from KeyManager.

    Returns:
        Tuple of (deidentified_text, entities_found).
    """
    if config.DRY_RUN:
        from mcp_deidentify.engine import SYNTHETIC_ENTITIES

        return _SYNTHETIC_DEID_NOTE, list(SYNTHETIC_ENTITIES)

    entities = await extract_entities(text)
    deidentified = replace_entities(text, entities, session_key, patient_id)
    return deidentified, entities


async def deidentify_docx_file(
    docx_path: str,
    patient_id: str,
    session_key: dict,
    output_path: str | None = None,
) -> tuple[str, str, list[dict]]:
    """De-identify a DOCX file and write a new de-identified DOCX.

    Processes all paragraph runs in parallel. Tables and images are preserved
    unchanged (text within table cells is NOT de-identified in Phase 3).

    Args:
        docx_path:   Path to the source DOCX file.
        patient_id:  Patient identifier for code generation.
        session_key: Mutable anonymization key dict from KeyManager.
        output_path: Destination path for de-identified DOCX.
                     Defaults to {OUTPUT_DIR}/{patient_id}/deidentified/{stem}_deid.docx

    Returns:
        Tuple of (deidentified_full_text, output_path_written, entities_found).
    """
    if output_path is None:
        output_path = _default_output_path(patient_id, docx_path)

    if config.DRY_RUN:
        from mcp_deidentify.engine import SYNTHETIC_ENTITIES

        return (
            _SYNTHETIC_DEID_NOTE,
            output_path + " [DRY_RUN — not written]",
            list(SYNTHETIC_ENTITIES),
        )

    from docx import Document

    doc = Document(docx_path)
    all_entities: list[dict] = []

    # Collect all paragraphs that have text
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    # De-identify all paragraphs in parallel
    tasks = [extract_entities(p.text) for p in paragraphs]
    results = await asyncio.gather(*tasks)

    full_text_parts = []
    for para, entities in zip(paragraphs, results):
        if entities:
            all_entities.extend(entities)
            new_text = replace_entities(para.text, entities, session_key, patient_id)
            # Clear existing runs and set de-identified text in the first run
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
        full_text_parts.append(para.text)

    # Write output
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info(f"De-identified DOCX written: {out}")

    return "\n".join(full_text_parts), output_path, all_entities
